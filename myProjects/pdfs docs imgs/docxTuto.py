import docx, os

# DOCX IS SOUP OF CLUSTERFUCK of  CONFUSING ELEMENTS< so TRY not to think and GO STRAIGHT FOR WHAT YOU WANT

#  Word Docx Tutorial
#these .paragraph, .cell .sheet functions store a list of touples where each touple represents a cells with its specific id( id and other shit is returned when print is given to function)

os.chdir('C:\\Users\\dhira\\Desktop\\mywork')
doc = docx.Document('demo.docx')    # everything from this document is imported eg, styles,fonts,etc It does not import shit from default styles so default styles are not present
                                    # if you want default styles keep doc.Document () blank
#print(len(doc.paragraphs))
#sprint((doc.paragraphs[0].text))  # returns actual text in given para
#print(len(doc.paragraphs[1].runs))
#for i in range(len(doc.paragraphs[1].runs)):
#    print((doc.paragraphs[1].runs[i].text))



def getText(filename):     # this function gets all the text from word document
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
        #fullText.append('     ' + para.text)    #indent each para
    return '\n\n'.join(fullText)
#print(getText('demo.docx'))





# STYLE Tuto
document = docx.Document()
paragraph = document.add_paragraph()

print(doc.paragraphs[0].style)
doc.paragraphs[0].style = 'Normal'   # changed style from Title style to Normal Style
print(doc.paragraphs[0].style)
print(doc.paragraphs[0].text)
#print(doc.paragraphs[1].runs[0].text, doc.paragraphs[1].runs[1].text, doc.paragraphs[1].runs[2].text, doc.paragraphs[1].runs[3].text)
print(doc.paragraphs[1].runs[0].text)
print(doc.paragraphs[1].runs[0].style.name)
doc.paragraphs[1].runs[0].style = document.styles['Heading1Char']  # OR JUST 'Heading1Char' OR JUST ONE OF RUN STYLES IN ACTUAL DOCUMENT eg ' Enphasis or Strong  # Changed Run style,  
print(doc.paragraphs[1].runs[0].style.name)                    # This will run an error telling to use style name isntead no matter what, ignore it. 
doc.paragraphs[1].runs[1].underline = True
doc.paragraphs[1].runs[3].underline = True
#print(doc.paragraphs[1].runs[0].style)


# THIS GETS ADDED TO NEXT PAGE
# WRITING TO WORD DOCS
doc.add_paragraph('Hellow World repeated', 'Heading2')   # style is optional   ( string_to_add , string_style)
paraObj1 = doc.add_paragraph('This is a second paragraph.')
paraObj2 = doc.add_paragraph('This is a yet another paragraph.')
paraObj1.add_run(' This text is being added to the second paragraph.')
doc.add_heading('Header 0', 0)   # ( String, 0-4 any size)

doc.add_paragraph('This is on the first page!')
print(doc.paragraphs[5].runs[0].text)
doc.paragraphs[5].runs[0].add_break
doc.add_paragraph('This is on the second page!')


#doc.add_picture('zophie.png', width=docx.shared.Inches(1), height=docx.shared.Cm(4))

doc.save('demo2.docx')

    
