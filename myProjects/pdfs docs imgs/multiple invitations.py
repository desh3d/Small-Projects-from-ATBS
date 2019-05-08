import docx,os,openpyxl

# This program creates a new docx file for each guest name(print all the files seperately or attach each to seperate mail), if you can, try combine all invitaions to one file

os.chdir('C:\\Users\\dhira\\Desktop\\mywork')
wbe = openpyxl.load_workbook('guests.xlsx')    # load an excel file containing names of all guests in first row
#print(wbe.sheetnames)  
sheet = wbe['Sheet1']
wbw = docx.Document('Invitation.docx')# load required sheet
for rowNum in range(2,sheet.max_row+1):
    guestname = sheet.cell(row=rowNum, column=1).value
    print(guestname)
    wbw.paragraphs[1].text = (guestname)
    print(wbw.paragraphs[1].text)
    #wbw.save(guestname+ ' invite'  + '.docx')    # create an individual invitation word file for each guest
    wbw.paragraphs[3].runs[2].add_page_break
    #wbw.paragraphs[paranum+2].runs[2].add_break(docx.enum.text.WD_BREAK.PAGE)   

#wbw.save('invitaions.docx')
    




#loop throught each para and run inside each para to get your desired para no and run no.
'''
wbw = docx.Document('Invitation.docx')

for i in range(len(wbw.paragraphs)):
    #print(wbw.paragraphs[i].text)
    print(i)
    for j in range(len(wbw.paragraphs[i].runs)):
        print(wbw.paragraphs[i].runs[j].text)

# to rep;ace and  save as a new docx document        
wbw.paragraphs[1].text = ('shoutdown', 'center') # cteae any style named 'center'  in the Invitation.docx and you can import and apply it dirctly to any char,or para like this.
print(wbw.paragraphs[1].text)
wbw.save('newinvite.docx')'''
